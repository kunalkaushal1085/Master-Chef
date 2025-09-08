import os, re, wave, numpy as np, pyaudio, threading, queue, time
from pathlib import Path
from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import UnstructuredWordDocumentLoader, UnstructuredPDFLoader
from langchain_chroma import Chroma
from langchain.memory import ConversationBufferWindowMemory
import random
import tempfile
import soundfile as sf
import pdf2image



POPPLER_PATH = r"C:/Users/Kunal Kaushal/Downloads/Release-25.07.0-0/poppler-25.07.0/Library\bin"

# Test if poppler is accessible
def verify_poppler():
    try:
        if os.path.exists(POPPLER_PATH):
            required_files = ['pdfinfo.exe', 'pdftoppm.exe']  # Only executables here!
            missing_files = [f for f in required_files if not os.path.exists(os.path.join(POPPLER_PATH, f))]
            if missing_files:
                print(f"Missing files in Poppler directory: {missing_files}")
            else:
                print("✓ Poppler path looks correct!")
        else:
            print(f"Poppler path does not exist: {POPPLER_PATH}")
    except Exception as e:
        print(f"Error checking Poppler: {e}")

# Your existing poppler path fix
_old_convert_from_path = pdf2image.convert_from_path

def _convert_from_path_fixed(*args, **kwargs):
    if "poppler_path" not in kwargs:
        kwargs["poppler_path"] = POPPLER_PATH
    return _old_convert_from_path(*args, **kwargs)

pdf2image.convert_from_path = _convert_from_path_fixed


# Enhanced principles extraction - Updated for Rosendale Method
KEEP = [
    r"philosophy", r"approach", r"why", r"because", r"ensures", r"prevents",
    r"results", r"should", r"never", r"always", r"critical", r"clarity",
    r"depth", r"restraint", r"fundamentals", r"technique", r"principle",
    r"coaching", r"guidance", r"mindset", r"understanding", r"mastery",
    r"teaches", r"learns", r"develops", r"builds", r"foundation",
    r"rules", r"control", r"texture", r"safety", r"avoid", r"use",
    r"essential", r"proper", r"method", r"ideal", r"recommended",
    r"precise", r"consistency", r"focus on", r"goal is",
    # Rosendale Method specific patterns
    r"rosendale method", r"key insight", r"common mistake", r"pro tip",
    r"what separates", r"the difference between", r"real secret",
    r"professional approach", r"kitchen wisdom", r"experience shows",
    r"problem solving", r"troubleshooting", r"when things go wrong",
    r"signature", r"technique breakdown", r"coaching moment"
]


EMOTIONAL_EXPRESSIONS = {
    "enthusiasm": [
        "I'm genuinely excited to share this principle with you!", 
        "This is where cooking mastery gets absolutely fascinating!",
        "Now we're diving into the real foundation of great cooking!",
        "This coaching moment is honestly one of my favorites!",
        "Here's where the Rosendale Method really shines!"
    ],
    "concern": [
        "I really want you to understand this crucial principle...",
        "This technique is critically important to master—",
        "Let me coach you through this because it truly matters—",
        "Please focus on this fundamental concept:",
        "Here's what genuinely concerns me when cooks skip this principle:"
    ],
    "pride": [
        "When you master this technique, you'll feel the difference immediately!",
        "There's honestly nothing like the confidence this principle builds...",
        "You'll know you've truly understood this when...",
        "This mastery is what really separates good cooks from exceptional ones!",
        "This is the foundation of professional cooking!"
    ],
    "empathy": [
        "I completely understand this principle can feel complex at first...",
        "Don't worry at all, every great cook has struggled with this concept—",
        "I remember feeling exactly the same way when learning this technique...",
        "It's totally normal to feel uncertain about this method...",
        "Every professional goes through this exact learning curve..."
    ],
    "curiosity": [
        "Have you ever genuinely wondered why this technique works...?",
        "What absolutely fascinates me about this method is...",
        "Think about this principle for just a moment—",
        "Here's something that might really surprise you about this technique:",
        "I'm genuinely curious—have you ever noticed this happening...?"
    ],
    "wisdom": [
        "Years in professional kitchens have taught me this principle...",
        "Here's what real experience has shown me about this technique:",
        "The deeper truth I've discovered through the Rosendale Method is...",
        "What I've learned through countless hours coaching is...",
        "After decades in the kitchen, this principle has proven essential..."
    ]
}


def extract_principles_from_pdf(src: str, dst: str):
    """Extract coaching principles from PDF documents"""
    if Path(dst).exists():
        return
    
    loader = UnstructuredPDFLoader(src)
    docs = loader.load()
    doc_out = Document()
    
    for doc in docs:
        text = doc.page_content
        paragraphs = text.split('\n')
        
        for para in paragraphs:
            para = para.strip()
            if para and any(re.search(pat, para, re.I) for pat in KEEP):
                para = re.sub(r'\s+', ' ', para)
                if len(para) > 20:
                    doc_out.add_paragraph(para)
    
    doc_out.save(dst)
    print(f"✓ PDF principles extracted → {dst}")


def recipe_to_principles(src: str, dst: str):
    """Extract coaching principles from Word documents"""
    if Path(dst).exists():
        return
    
    if not Path(src).exists():
        print(f"Warning: {src} not found, skipping...")
        return
        
    doc_in, doc_out = Document(src), Document()
    for p in doc_in.paragraphs:
        t = p.text.strip()
        if t and any(re.search(pat, t, re.I) for pat in KEEP):
            doc_out.add_paragraph(t)
    doc_out.save(dst)
    print(f"✓ Word doc principles extracted → {dst}")


def get_emotional_opener(context_keywords):
    """Enhanced emotional context detection for coaching"""
    context_lower = context_keywords.lower()
    
    if any(word in context_lower for word in ["safety", "dangerous", "avoid", "never", "mistake"]):
        return random.choice(EMOTIONAL_EXPRESSIONS["concern"]), "concern"
    elif any(word in context_lower for word in ["master", "perfect", "technique", "skill", "professional"]):
        return random.choice(EMOTIONAL_EXPRESSIONS["pride"]), "pride"
    elif any(word in context_lower for word in ["why", "science", "understand", "fascinating", "principle"]):
        return random.choice(EMOTIONAL_EXPRESSIONS["curiosity"]), "curiosity"
    elif any(word in context_lower for word in ["difficult", "hard", "struggle", "problem", "troubleshoot"]):
        return random.choice(EMOTIONAL_EXPRESSIONS["empathy"]), "empathy"
    elif any(word in context_lower for word in ["experience", "learned", "years", "wisdom", "rosendale"]):
        return random.choice(EMOTIONAL_EXPRESSIONS["wisdom"]), "wisdom"
    else:
        return random.choice(EMOTIONAL_EXPRESSIONS["enthusiasm"]), "enthusiasm"


def init_chain():
    """Initialize the Rosendale Method coaching chain"""
    load_dotenv()
    openai_client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

    # Extract principles from available documents
    recipe_to_principles("48 Hour Short Ribs.docx", "Rosendale_Principles.docx")
    extract_principles_from_pdf("Sous-Vide-Fish-1.pdf", "Fish_Principles.docx")
    
    all_docs = []
    
    # Load available principle documents
    for doc_path in ["Rosendale_Principles.docx", "Fish_Principles.docx"]:
        if Path(doc_path).exists():
            docs = UnstructuredWordDocumentLoader(doc_path).load()
            all_docs.extend(docs)
            print(f"✓ Loaded {len(docs)} documents from {doc_path}")
    
    if not all_docs:
        print("Warning: No principle documents found!")
        return openai_client, None, None, None
    
    # Optimized chunking for coaching principles
    splitter = RecursiveCharacterTextSplitter(chunk_size=400, chunk_overlap=80)
    split_docs = splitter.split_documents(all_docs)
    print(f"✓ Split into {len(split_docs)} principle chunks")

    # Create vector store for principle retrieval
    vectordb = Chroma.from_documents(
        split_docs, 
        OpenAIEmbeddings(), 
        persist_directory="./chroma_rosendale"
    )   
    retriever = vectordb.as_retriever(search_kwargs={"k": 3})


    # Conversation memory for coaching context
    memory = ConversationBufferWindowMemory(
        memory_key="chat_history",
        return_messages=True, 
        k=3  # Keep recent context for coaching flow
    )
    
    # Optimized LLM for coaching responses
    llm = ChatOpenAI(
        model="gpt-4o", 
        temperature=0.7, 
        max_tokens=350  # Concise coaching responses
    )
    
    print("✓ Rosendale Method coaching chain initialized")
    return openai_client, retriever, memory, llm


def mentor_answer(q, retriever, memory, llm):
    """Generate detailed conversational coaching responses using Rosendale Method"""
    if not retriever:
        return ("I'm still getting familiar with your cooking principles, but I'm excited to start coaching you! Every great cook begins with curiosity about technique. What specific cooking challenge would you like to explore together?", "empathy")
    
    # Use the new invoke method
    ctx_blocks = retriever.invoke(q)
    if not ctx_blocks:
        response = ("I love diving into cooking principles with passionate cooks like yourself! "
                   "There's so much wisdom to share about technique, timing, and the science behind great food. "
                   "What culinary challenge has been on your mind lately? I'm here to help you master it.")
        return response, "enthusiasm"
    
    # Prepare context from principles
    context = "\n".join(d.page_content for d in ctx_blocks[:3])[:1800]
    
    # Get conversation history
    history_messages = memory.chat_memory.messages[-6:]
    history = "\n".join(f"{m.type.upper()}: {m.content}" for m in history_messages)[:800]
    
    # Get emotion type without the opener text
    _, emotion_type = get_emotional_opener(context + " " + q)
    
    # Enhanced coaching prompt WITHOUT emotional opener injection
    prompt = (
        f"COACHING CONTEXT: {context}\n"
        f"CONVERSATION HISTORY: {history}\n"
        f"STUDENT QUESTION: {q}\n\n"
        
        "You are an experienced culinary coach using the Rosendale Method. "
        "Your responses should be conversational, warm, and detailed enough to truly help students understand. "
        
        "RESPONSE GUIDELINES:\n"
        "- Be conversational and engaging, like talking to a friend in the kitchen\n"
        "- Explain WHY techniques work with genuine enthusiasm and detail\n"
        "- Share professional insights and problem-solving wisdom\n"
        "- Use specific examples to illustrate principles, but teach concepts, not procedures\n"
        "- Show your passion for cooking through your explanations\n"
        "- Make it feel like a mentoring conversation, not a lecture\n"
        "- Length: 4-6 sentences that flow naturally together\n"
        "- Start directly with your teaching point - avoid filler words like 'Ah,' 'Oh,' etc.\n"
        "- End with an engaging question that encourages deeper exploration\n"
        "- Never provide step-by-step recipes or mention specific chef names\n"
        "- Focus on building understanding and confidence\n"
        
        "Remember: You're building a cook's intuition through clear, direct teaching."
    )
    
    try:
        response = llm.invoke(prompt).content.strip()
        
        # Clean up any remaining filler expressions
        response = re.sub(r"^(Ah+|Oh+|Um+|Well+),?\s*", '', response, flags=re.I)
        response = re.sub(r"^(You know|Listen|Look),?\s*", '', response, flags=re.I)
        
        # Ensure proper ending punctuation
        if not response.endswith(('.', '!', '?')):
            response += '.'
        
        # Update conversation memory
        memory.chat_memory.add_user_message(q)
        memory.chat_memory.add_ai_message(response)
        
        return response, emotion_type
        
    except Exception as e:
        print(f"Error in mentor_answer: {e}")
        return ("That's a fascinating question that really gets to the heart of cooking mastery. Let me think about the best way to break this down for you. What specific aspect of this technique has been challenging you the most?", "empathy")


def is_speech(audio_data, threshold=0.005):
    """Check if audio contains speech"""
    rms = np.sqrt(np.mean(np.square(audio_data)))
    return rms > threshold


def is_valid_transcription(text):
    """Validate transcription quality"""
    meaningless = {
        "thank you", "thank you for watching", "hello", "hi", "...", "", "welcome back",
        "thank you everyone", "thanks for watching", "okay", "ok", "yes", "no", "nothing",
        "Thank you for watching!", "Thank you for watching", " Thank you for watching."
    }
    cleaned = text.strip().lower()
    non_latin_ratio = sum(1 for c in cleaned if not re.match(r'[a-zA-Z0-9\s,.!?]', c)) / max(len(cleaned), 1)
    return (
        len(cleaned) > 5 and
        cleaned not in meaningless and
        non_latin_ratio < 0.3
    )


def transcribe_audio(audio_bytes, client):
    """Transcribe audio to text using OpenAI Whisper"""
    # Create temporary file for audio
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
        temp_audio.write(audio_bytes)
        audio_path = temp_audio.name

    try:
        with open(audio_path, "rb") as audio_file:
            transcription = client.audio.transcriptions.create(
                model="whisper-1",  # Fixed model name
                file=audio_file,
                response_format="text",
                language="en",
                temperature=0.2,
                prompt="Listen for cooking and culinary questions about techniques, methods, and principles."
            )
        os.remove(audio_path)
        
        if not is_valid_transcription(transcription):
            return None
        return transcription
        
    except Exception as e:
        print("Transcription error:", e)
        if os.path.exists(audio_path):
            os.remove(audio_path)
        return None


def autocomplete_if_needed(user_input, memory, llm):
    """Enhance short queries for better coaching context"""
    if len(user_input.split()) < 4:
        history = memory.chat_memory.messages[-6:]
        chat_history_str = "\n".join(f"{m.type.upper()}: {m.content}" for m in history)
        prompt = (
            f"The student asked: '{user_input}'\n"
            f"Coaching conversation so far:\n{chat_history_str}\n"
            f"Assume they're asking about cooking techniques or culinary principles. "
            f"What is the most likely complete question they meant to ask about cooking mastery?"
        )
        try:
            response = llm.invoke(prompt)
            return response.content.strip()
        except:
            return user_input
    return user_input


def is_recipe_related(query):
    """Check if query is related to cooking/culinary topics"""
    cooking_keywords = [
        "recipe", "cook", "bake", "grill", "roast", "sauté", "fry", "steam", "boil",
        "sauce", "garlic", "onion", "burn", "temperature", "heat", "pan", "oil",
        "technique", "method", "flavor", "texture", "ingredient", "season", "salt"
    ]
    
    # Also check for common cooking question patterns
    cooking_patterns = [
        r"why does.*burn", r"how do.*cook", r"what.*temperature",
        r"how.*prevent", r"why.*tough", r"how.*make", r"why.*dry"
    ]
    
    query_lower = query.lower()
    
    # Check keywords
    if any(word in query_lower for word in cooking_keywords):
        return True
    
    # Check patterns
    if any(re.search(pattern, query_lower) for pattern in cooking_patterns):
        return True
        
    return False


def get_qa_answer(full_query, qa_chain):
    """Legacy function for backward compatibility"""
    response = qa_chain.invoke({"question": full_query})
    return response["answer"]



#for testing purpose this function questions and answer
# import pygame
# import tempfile
# import os
# from elevenlabs_functions import speak_text_to_stream
# import speech_recognition as sr

# def terminal_voice_chat():
#     """
#     Terminal-based voice chat function with speech input
#     - SPEAK your question (no typing)
#     - Get AI response 
#     - Speak response using ElevenLabs voice
#     """
    
#     print("\n🍳 Master Chef Voice-to-Voice Chat")
#     print("=" * 50)
#     print("🗣️ SPEAK your cooking questions - no typing needed!")
#     print("🛑 Say 'quit', 'exit', or 'stop' to end")
#     print("=" * 50)
    
#     # Initialize components
#     try:
#         client, retriever, memory, llm = init_chain()
#         if not retriever:
#             print("❌ Failed to initialize cooking AI. Please check your setup.")
#             return
#     except Exception as e:
#         print(f"❌ Initialization error: {e}")
#         return
    
#     # Initialize speech recognition and audio playback
#     try:
#         recognizer = sr.Recognizer()
#         microphone = sr.Microphone()
#         pygame.mixer.init()
#         print("✅ Voice recognition and audio system ready")
#     except Exception as e:
#         print(f"❌ Audio initialization failed: {e}")
#         return
    
#     # Calibrate microphone
#     print("🎤 Calibrating microphone... please wait")
#     with microphone as source:
#         recognizer.adjust_for_ambient_noise(source, duration=1)
#     print("✅ Microphone calibrated!")
    
#     while True:
#         try:
#             # LISTEN FOR QUESTION (replaces input())
#             print("\n" + "-" * 60)
#             print("🎤 Speak your cooking question now...")
            
#             with microphone as source:
#                 audio = recognizer.listen(source, timeout=5, phrase_time_limit=10)
            
#             try:
#                 # Convert speech to text
#                 question = recognizer.recognize_google(audio)
#                 print(f"🎯 You asked: {question}")
                
#                 # Handle quit commands
#                 if any(word in question.lower() for word in ['quit', 'exit', 'stop']):
#                     print("👋 Happy cooking! Goodbye!")
#                     break
                
#                 # Skip empty questions
#                 if not question.strip():
#                     print("⚠️  Please ask a cooking question")
#                     continue
                
#                 print(f"🔍 Processing: {question}")
                
#                 # Enhance short questions
#                 full_query = autocomplete_if_needed(question, memory, llm)
#                 if full_query != question:
#                     print(f"🎯 Enhanced to: {full_query}")
                
#                 # Validate cooking-related content
#                 if not is_recipe_related(full_query):
#                     print("❌ Please ask about cooking techniques, methods, or culinary principles.")
#                     continue
                
#                 # Get AI response
#                 print("🤖 Chef is thinking...")
#                 answer, emotion_type = mentor_answer(full_query, retriever, memory, llm)
                
#                 # Display text response
#                 print("\n📋 CHEF'S RESPONSE:")
#                 print("-" * 40)
#                 print(answer)
#                 print("-" * 40)
#                 print(f"😊 Response emotion: {emotion_type}")
                
#                 # Convert to voice
#                 print("\n🎤 Converting to voice...")
#                 audio_stream = speak_text_to_stream(answer)
#                 audio_bytes = audio_stream.getvalue()
                
#                 if len(audio_bytes) == 0:
#                     print("❌ Could not generate voice response")
#                     continue
                
#                 # Play voice response
#                 print("🔊 Playing voice response...")
                
#                 # Save to temporary file and play
#                 with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp_file:
#                     temp_file.write(audio_bytes)
#                     temp_audio_path = temp_file.name
                
#                 try:
#                     # Play audio
#                     pygame.mixer.music.load(temp_audio_path)
#                     pygame.mixer.music.play()
                    
#                     print("🎵 Speaking... (wait for completion)")
                    
#                     # Wait for playback to finish
#                     while pygame.mixer.music.get_busy():
#                         pygame.time.wait(100)
                    
#                     print("✅ Voice response completed")
                    
#                 except Exception as e:
#                     print(f"❌ Audio playback error: {e}")
#                 finally:
#                     # Clean up temp file
#                     try:
#                         os.unlink(temp_audio_path)
#                     except:
#                         pass
                
#             except sr.UnknownValueError:
#                 print("🤔 Sorry, I couldn't understand that. Please speak more clearly.")
#             except sr.RequestError as e:
#                 print(f"❌ Speech recognition error: {e}")
                
#         except KeyboardInterrupt:
#             print("\n👋 Session ended. Happy cooking!")
#             break
#         except Exception as e:
#             print(f"❌ Error: {str(e)}")
#             print("🔄 Please try again...")
    
#     # Cleanup
#     pygame.mixer.quit()

# def voice_input_question():
#     """
#     Get ONE question from voice input and return it as text
#     (Replaces the input() function)
#     """
#     try:
#         recognizer = sr.Recognizer()
#         microphone = sr.Microphone()
        
#         # Quick calibration
#         with microphone as source:
#             recognizer.adjust_for_ambient_noise(source, duration=0.5)
        
#         print("🎤 Speak your question now...")
#         with microphone as source:
#             audio = recognizer.listen(source, timeout=5, phrase_time_limit=10)
        
#         question = recognizer.recognize_google(audio)
#         print(f"🎯 You said: {question}")
#         return question
        
#     except sr.UnknownValueError:
#         print("❌ Could not understand. Please try again.")
#         return None
#     except sr.RequestError as e:
#         print(f"❌ Speech recognition error: {e}")
#         return None
#     except Exception as e:
#         print(f"❌ Error: {e}")
#         return None
# terminal_voice_chat()
# # question = input("🥘 Ask your cooking question: ").strip()
# with microphone as source:
#     audio = recognizer.listen(source, timeout=5, phrase_time_limit=10)
# question = recognizer.recognize_google(audio)


import pygame
import tempfile
import os
import speech_recognition as sr
import time
from elevenlabs_functions import speak_text_to_stream

class MasterChefVoiceChat:
    """Professional voice-to-voice cooking assistant with proper exit detection"""
    
    def __init__(self):
        self.recognizer = sr.Recognizer()
        self.microphone = None
        self.client = None
        self.retriever = None
        self.memory = None
        self.llm = None
        self.session_active = False
        
        # Configure audio parameters for better recognition
        self.setup_audio_parameters()
    
    def setup_audio_parameters(self):
        """Configure optimal audio parameters for clear recognition"""
        self.recognizer.energy_threshold = 300
        self.recognizer.dynamic_energy_threshold = True
        self.recognizer.dynamic_energy_adjustment_damping = 0.15
        self.recognizer.dynamic_energy_ratio = 1.5
        self.recognizer.pause_threshold = 0.8
        self.recognizer.phrase_threshold = 0.3
        self.recognizer.non_speaking_duration = 0.5
        
    def is_exit_command(self, text):
        """
        IMPROVED exit detection - stops immediately when user says goodbye phrases
        """
        if not text or len(text.strip()) < 1:
            return False
            
        text_lower = text.lower().strip()
        
        # Primary exit phrases - immediate stop
        primary_exit_phrases = [
            'quit', 'exit', 'stop', 'bye', 'goodbye', 'good bye',
            'thank you', 'thanks', 'end', 'finish', 'done'
        ]
        
        # Check for exact matches or phrases at start/end
        for phrase in primary_exit_phrases:
            if phrase == text_lower:  # Exact match
                print(f"✓ Exit phrase detected: '{phrase}'")
                return True
            elif text_lower.startswith(phrase + ' '):  # Starts with phrase
                print(f"✓ Exit phrase detected at start: '{phrase}'")
                return True
            elif text_lower.endswith(' ' + phrase):  # Ends with phrase
                print(f"✓ Exit phrase detected at end: '{phrase}'")
                return True
            elif phrase in text_lower and len(text_lower.split()) <= 4:  # Short phrase containing exit word
                print(f"✓ Exit phrase detected in short sentence: '{phrase}'")
                return True
        
        # Special patterns that indicate end of conversation
        goodbye_patterns = [
            'thank you goodbye', 'thanks goodbye', 'thank you bye',
            'thanks bye', 'okay bye', 'ok bye', 'okay thank you',
            'ok thanks', 'alright bye', 'that\'s all', 'no more',
            'i\'m done', 'all done', 'stop now', 'end now'
        ]
        
        for pattern in goodbye_patterns:
            if pattern in text_lower:
                print(f"✓ Goodbye pattern detected: '{pattern}'")
                return True
        
        return False
    
    def initialize_systems(self):
        """Initialize all required systems"""
        print("\n🍳 Master Chef Professional Voice Assistant")
        print("=" * 60)
        
        try:
            # Initialize AI components
            print("🤖 Initializing AI cooking coach...")
            self.client, self.retriever, self.memory, self.llm = init_chain()
            if not self.retriever:
                print("❌ Failed to initialize cooking AI.")
                return False
            
            # Initialize audio systems
            print("🎤 Setting up audio systems...")
            self.microphone = sr.Microphone()
            pygame.mixer.init(frequency=22050, size=-16, channels=2, buffer=1024)
            
            # Calibrate microphone
            self.calibrate_microphone()
            
            print("✅ All systems ready!")
            print("\n🗣️  VOICE CONVERSATION MODE")
            print("🎯 Speak naturally about cooking")
            print("👋 Say 'thank you', 'bye', or 'quit' to end")
            print("=" * 60)
            
            self.session_active = True
            return True
            
        except Exception as e:
            print(f"❌ Initialization failed: {e}")
            return False
    
    def calibrate_microphone(self):
        """Calibrate microphone for optimal performance"""
        print("🎤 Calibrating microphone...")
        with self.microphone as source:
            self.recognizer.adjust_for_ambient_noise(source, duration=1.5)
        print(f"✅ Calibration complete (threshold: {int(self.recognizer.energy_threshold)})")
    
    def listen_for_question(self):
        """Listen for voice input with retry logic"""
        max_retries = 2
        
        for attempt in range(max_retries):
            try:
                print(f"\n👂 Listening... (attempt {attempt + 1}/{max_retries})")
                print("🗣️  Speak clearly now...")
                
                with self.microphone as source:
                    audio = self.recognizer.listen(
                        source, 
                        timeout=6,
                        phrase_time_limit=12
                    )
                
                print("🔄 Processing speech...")
                
                question = self.recognizer.recognize_google(audio, language="en-US")
                
                if question and len(question.strip()) > 0:
                    return question.strip()
                    
            except sr.WaitTimeoutError:
                if attempt < max_retries - 1:
                    print("⏰ No speech detected. Please try again.")
                else:
                    print("❌ No speech detected after multiple attempts.")
                    
            except sr.UnknownValueError:
                if attempt < max_retries - 1:
                    print("🤔 Couldn't understand clearly. Please repeat.")
                else:
                    print("❌ Unable to understand speech.")
                    
            except sr.RequestError as e:
                print(f"❌ Speech service error: {e}")
                return None
                
            except Exception as e:
                print(f"❌ Unexpected error: {e}")
        
        return None
    
    def speak_response(self, text, context="response"):
        """Convert text to speech and play"""
        try:
            if context == "greeting":
                print("👋 Playing welcome message...")
            elif context == "goodbye":
                print("👋 Playing farewell message...")
            else:
                print("🎤 Converting to speech...")
                
            audio_stream = speak_text_to_stream(text)
            audio_bytes = audio_stream.getvalue()
            
            if len(audio_bytes) == 0:
                print("❌ Voice generation failed")
                return False
            
            print("🔊 Speaking...")
            
            with tempfile.NamedTemporaryFile(delete=False, suffix='.mp3') as temp_file:
                temp_file.write(audio_bytes)
                temp_audio_path = temp_file.name
            
            try:
                pygame.mixer.music.load(temp_audio_path)
                pygame.mixer.music.play()
                
                while pygame.mixer.music.get_busy():
                    time.sleep(0.1)
                
                print("✅ Speech completed")
                return True
                
            finally:
                try:
                    os.unlink(temp_audio_path)
                except:
                    pass
                    
        except Exception as e:
            print(f"❌ Speech error: {e}")
            return False
    
    def process_cooking_question(self, question):
        """Process cooking question and generate response"""
        try:
            print(f"🎯 Processing: \"{question}\"")
            
            # Enhance question if needed
            full_query = autocomplete_if_needed(question, self.memory, self.llm)
            if full_query != question:
                print(f"🎯 Enhanced: \"{full_query}\"")
            
            # Validate cooking topic
            if not is_recipe_related(full_query):
                response = "I specialize in cooking techniques and culinary principles. What cooking topic would you like to explore?"
                print("❌ Not cooking-related - redirecting")
                self.speak_response(response)
                return True
            
            # Generate AI response
            print("🤖 Chef thinking...")
            answer, emotion_type = mentor_answer(full_query, self.retriever, self.memory, self.llm)
            
            # Show response
            print(f"\n📋 CHEF'S ADVICE:")
            print("-" * 50)
            print(answer)
            print("-" * 50)
            print(f"😊 Tone: {emotion_type}")
            
            # Speak response
            success = self.speak_response(answer)
            
            if success:
                print("✅ Response delivered")
            
            return True
            
        except Exception as e:
            print(f"❌ Processing error: {e}")
            error_msg = "I had trouble with that question. Could you please try again?"
            self.speak_response(error_msg)
            return True
    
    def run_conversation(self):
        """Main conversation loop with proper exit handling"""
        if not self.initialize_systems():
            return
        
        # Welcome message
        welcome = "Hello! I'm your Master Chef voice assistant. What cooking technique would you like to learn today?"
        self.speak_response(welcome, "greeting")
        
        conversation_count = 0
        
        while self.session_active:
            try:
                # Listen for user input
                question = self.listen_for_question()
                
                if question is None:
                    print("🔄 Let's try listening again...")
                    continue
                
                # **CRITICAL FIX**: Check for exit IMMEDIATELY
                if self.is_exit_command(question):
                    print(f"🛑 EXIT DETECTED: \"{question}\"")
                    
                    # Speak farewell and end immediately
                    farewell = "Thank you for cooking with me today! Keep practicing those techniques. Happy cooking!"
                    self.speak_response(farewell, "goodbye")
                    
                    print("👋 Session ended by user")
                    break  # **IMMEDIATELY EXIT THE LOOP**
                
                # Process as cooking question only if NOT an exit command
                conversation_count += 1
                print(f"\n🗣️  Conversation #{conversation_count}")
                
                self.process_cooking_question(question)
                
                # Brief pause between interactions
                time.sleep(0.3)
                
            except KeyboardInterrupt:
                print("\n👋 Session interrupted")
                break
            except Exception as e:
                print(f"❌ Error: {e}")
                continue
        
        # Final cleanup
        self.cleanup()
    
    def cleanup(self):
        """Clean shutdown"""
        print("\n🧹 Cleaning up...")
        try:
            pygame.mixer.quit()
            print("✅ Audio systems closed")
        except:
            pass
        print("👨‍🍳 Master Chef session complete")

# Simple function to start the chat
def start_voice_chat():
    """Start the voice conversation"""
    chef = MasterChefVoiceChat()
    chef.run_conversation()

# Run it
if __name__ == "__main__":
    start_voice_chat()

